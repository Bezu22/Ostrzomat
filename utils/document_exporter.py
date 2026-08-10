import os
import sqlite3
import hashlib
from datetime import datetime
from typing import Optional

# Łątka dla użycia md5 w bibliotekach
_original_md5 = hashlib.md5
def _safe_md5(*args, **kwargs):
    kwargs.pop('usedforsecurity', None)
    return _original_md5(*args, **kwargs)
hashlib.md5 = _safe_md5

from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def register_custom_fonts():
    try:
        fonts_dir = os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts')
        arial_path = os.path.join(fonts_dir, 'arial.ttf')
        arial_bold_path = os.path.join(fonts_dir, 'arialbd.ttf')

        pdfmetrics.registerFont(TTFont('ArialCustom', arial_path))
        pdfmetrics.registerFont(TTFont('ArialCustom-Bold', arial_bold_path))
        return 'ArialCustom', 'ArialCustom-Bold'
    except Exception:
        return 'Helvetica', 'Helvetica-Bold'


def fetch_client_data(client_id: Optional[int], db_path: Optional[str] = None) -> dict:
    if not client_id:
        return {
            "id": 0,
            "name": "Nieokreślony klient",
            "phone": "-",
            "nip": "-",
            "email": "-",
            "address": "-"
        }

    if db_path is None:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        db_path = os.path.join(os.path.dirname(script_dir), "data", "clients.db")

    if not os.path.exists(db_path):
        return {
            "id": client_id,
            "name": "Klient (Brak bazy)",
            "phone": "-", "nip": "-", "email": "-", "address": "-"
        }

    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("SELECT id, name, phone, nip, email, address FROM clients WHERE id = ?", (client_id,))
    row = cursor.fetchone()
    conn.close()

    if not row:
        return {
            "id": client_id,
            "name": "Nieznany klient",
            "phone": "-", "nip": "-", "email": "-", "address": "-"
        }

    return {
        "id": row[0],
        "name": row[1] or "Brak nazwy",
        "phone": row[2] or "-",
        "nip": row[3] or "-",
        "email": row[4] or "-",
        "address": row[5] or "-"
    }


def generate_pdf(cart_data: dict, client_info: dict, output_pdf_path: str):
    font_normal, font_bold = register_custom_fonts()

    doc = SimpleDocTemplate(
        output_pdf_path,
        pagesize=landscape(A4),
        rightMargin=20, leftMargin=20, topMargin=20, bottomMargin=20
    )

    story = []
    styles = getSampleStyleSheet()

    style_title = ParagraphStyle('DocTitle', parent=styles['Normal'], fontName=font_bold, fontSize=14, textColor=colors.HexColor('#1e3a8a'))
    style_normal = ParagraphStyle('DocNormal', parent=styles['Normal'], fontName=font_normal, fontSize=8, leading=10)
    style_cell = ParagraphStyle('CellText', parent=styles['Normal'], fontName=font_normal, fontSize=7, leading=8)
    style_cell_header = ParagraphStyle('CellHeader', parent=styles['Normal'], fontName=font_bold, fontSize=7, leading=8, textColor=colors.white, alignment=1)

    # --- NAGŁÓWEK DOKUMENTU ---
    header_data = [
        [
            Paragraph("<b>CENTRALA TECHNICZNA</b><br/><font size=8 color='#475569'>REGENERACJA NARZĘDZI</font>", style_title),
            Paragraph(f"<para align='right'><b>WYCENA ZLECENIA</b><br/>Data: {datetime.now().strftime('%d.%m.%Y r.')}</para>", style_normal)
        ]
    ]
    header_table = Table(header_data, colWidths=[400, 400])
    header_table.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP')]))
    story.append(header_table)
    story.append(Spacer(1, 10))

    # --- DANE KLIENTA ---
    client_text = (
        f"<b>DANE KLIENTA:</b> {client_info['name']} &nbsp;&nbsp;|&nbsp;&nbsp; "
        f"<b>NIP:</b> {client_info['nip']}<br/>"
        f"<b>Adres:</b> {client_info['address']} &nbsp;&nbsp;|&nbsp;&nbsp; "
        f"<b>Tel:</b> {client_info['phone']} &nbsp;&nbsp;|&nbsp;&nbsp; "
        f"<b>Email:</b> {client_info['email']}"
    )
    client_table = Table([[Paragraph(client_text, style_normal)]], colWidths=[800])
    client_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f8fafc')),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#cbd5e1')),
        ('PADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(client_table)
    story.append(Spacer(1, 10))

    # --- NAGŁÓWKI TABELI ---
    headers = [
        "Lp.", "Typ narzędzia", "Ø narz.", "Ø trz.", "L [mm]", "Ostrza", 
        "Ilość", "Cena ostrz.", "Wartość os.", "Powłoka", "Cena powł.", "Wartość powł.", "Usługi", "Wartość usł.", "Uwagi"
    ]
    
    table_data = [[Paragraph(h, style_cell_header) for h in headers]]

    items = cart_data.get("items", [])
    total_qty = 0
    total_tool_val = 0.0
    total_coat_val = 0.0
    total_extra_val = 0.0

    for idx, item in enumerate(items, start=1):
        qty = int(item.get("qty", 0))
        t_tool = float(item.get("total_tool", 0.0))
        t_coat = float(item.get("total_coat", 0.0))
        t_extra = float(item.get("total_extra", 0.0))

        total_qty += qty
        total_tool_val += t_tool
        total_coat_val += t_coat
        total_extra_val += t_extra

        st = item.get("services_status", {})
        services_list = []
        if st.get("ciecie"): services_list.append("Cięcie")
        if st.get("opuszczenie"): services_list.append("Opuszczenie")
        if st.get("polerowanie"): services_list.append("Polerowanie")
        if st.get("zuzycie"): services_list.append("Zużycie")
        uslugi_str = ", ".join(services_list) if services_list else "-"

        row = [
            Paragraph(str(idx), style_cell),
            Paragraph(f"<b>{item.get('type', '-')}</b>", style_cell),
            Paragraph(f"Ø{item.get('diam', '-')}", style_cell),
            Paragraph(f"Ø{item.get('shank_diam', '-')}", style_cell),
            Paragraph(str(item.get('coat_len', '-')), style_cell),
            Paragraph(str(item.get('z', '-')), style_cell),
            Paragraph(f"<b>{qty}</b>", style_cell),
            Paragraph(f"{float(item.get('tool_unit', 0)):.2f} zł", style_cell),
            Paragraph(f"{t_tool:.2f} zł", style_cell),
            Paragraph(str(item.get('coat_name', '-')), style_cell),
            Paragraph(f"{float(item.get('coat_unit', 0)):.2f} zł", style_cell),
            Paragraph(f"{t_coat:.2f} zł", style_cell),
            Paragraph(uslugi_str, style_cell),
            Paragraph(f"{t_extra:.2f} zł", style_cell),
            Paragraph(str(item.get('notes', '')), style_cell)
        ]
        table_data.append(row)

    col_widths = [25, 95, 35, 35, 40, 30, 30, 50, 55, 80, 50, 55, 70, 55, 145]
    items_table = Table(table_data, colWidths=col_widths, repeatRows=1)
    
    # BEDZIE STOSOWANE BEZPOŚREDNIE MAPOWANIE TŁA BEZ KONFLIKTU STYLÓW
    ts = TableStyle()
    # 1. Tło wierszy danych (od wiersza 1 do końca): biały oraz szary (#cbd5e1)
    ts.add('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#cbd5e1')])
    # 2. Granatowe tło nagłówka (wiersz 0) nadpisuje tło z ROWBACKGROUNDS
    ts.add('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e3a8a'))
    # 3. Siatka i wyrównanie
    ts.add('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#94a3b8'))
    ts.add('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
    ts.add('TOPPADDING', (0, 0), (-1, -1), 3)
    ts.add('BOTTOMPADDING', (0, 0), (-1, -1), 3)

    items_table.setStyle(ts)
    story.append(items_table)
    story.append(Spacer(1, 10))

    # --- PODSUMOWANIE ---
    grand_total = total_tool_val + total_coat_val + total_extra_val
    summary_text = f"""
    <b>Suma sztuk:</b> {total_qty} szt.<br/>
    Wartość regeneracji: {total_tool_val:.2f} zł<br/>
    Wartość powlekania: {total_coat_val:.2f} zł<br/>
    Wartość usług dod.: {total_extra_val:.2f} zł<br/>
    <font size=10 color='#1e3a8a'><b>RAZEM NETTO: {grand_total:.2f} zł</b></font>
    """
    summary_table = Table([[Paragraph(summary_text, style_normal)]], colWidths=[250])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f1f5f9')),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#cbd5e1')),
        ('PADDING', (0,0), (-1,-1), 8),
        ('ALIGN', (0,0), (-1,-1), 'RIGHT')
    ]))
    
    layout_summary = Table([["", summary_table]], colWidths=[550, 250])
    story.append(layout_summary)

    doc.build(story)


def set_cell_background(cell, fill_hex: str):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def generate_docx(cart_data: dict, client_info: dict, output_docx_path: str):
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Inches(6.0)
    header_table.columns[1].width = Inches(4.5)

    cell_left = header_table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    run_title = p_left.add_run("CENTRALA TECHNICZNA\n")
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    run_sub = p_left.add_run("REGENERACJA NARZĘDZI")
    run_sub.font.size = Pt(8)
    run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    cell_right = header_table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_doc_title = p_right.add_run("WYCENA ZLECENIA\n")
    run_doc_title.bold = True
    run_doc_title.font.size = Pt(14)
    
    run_date = p_right.add_run(f"Data: {datetime.now().strftime('%d.%m.%Y r.')}")
    run_date.font.size = Pt(8)

    doc.add_paragraph()

    client_table = doc.add_table(rows=1, cols=1)
    client_cell = client_table.cell(0, 0)
    set_cell_background(client_cell, "F8FAFC")

    p_client = client_cell.paragraphs[0]
    p_client.paragraph_format.space_after = Pt(2)
    p_client.paragraph_format.space_before = Pt(2)
    
    client_str = (
        f"DANE KLIENTA: {client_info['name']} (ID: {client_info['id']})   |   NIP: {client_info['nip']}\n"
        f"Adres: {client_info['address']}   |   Tel: {client_info['phone']}   |   Email: {client_info['email']}"
    )
    run_client = p_client.add_run(client_str)
    run_client.font.size = Pt(8.5)
    run_client.font.name = 'Arial'

    doc.add_paragraph()

    headers = [
        "Lp.", "Typ narzędzia", "Ø narz.", "Ø trz.", "L [mm]", "Ostrza", 
        "Ilość", "Cena ostrz.", "Wartość", "Powłoka", "Cena powł.", "Wartość", "Usługi", "Wartość dod.", "Uwagi"
    ]

    items = cart_data.get("items", [])
    items_table = doc.add_table(rows=len(items) + 1, cols=15)
    items_table.style = 'Table Grid'

    hdr_cells = items_table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E3A8A")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = 'Arial'

    total_qty = 0
    total_tool_val = 0.0
    total_coat_val = 0.0
    total_extra_val = 0.0

    for idx, item in enumerate(items, start=1):
        row_cells = items_table.rows[idx].cells
        
        qty = int(item.get("qty", 0))
        t_tool = float(item.get("total_tool", 0.0))
        t_coat = float(item.get("total_coat", 0.0))
        t_extra = float(item.get("total_extra", 0.0))

        total_qty += qty
        total_tool_val += t_tool
        total_coat_val += t_coat
        total_extra_val += t_extra

        st = item.get("services_status", {})
        services_list = []
        if st.get("ciecie"): services_list.append("Cięcie")
        if st.get("opuszczenie"): services_list.append("Opuszczenie")
        if st.get("polerowanie"): services_list.append("Polerowanie")
        if st.get("zuzycie"): services_list.append("Zużycie")
        uslugi_str = ", ".join(services_list) if services_list else "-"

        row_data = [
            str(idx),
            str(item.get('type', '-')),
            f"Ø{item.get('diam', '-')}",
            f"Ø{item.get('shank_diam', '-')}",
            str(item.get('coat_len', '-')),
            str(item.get('z', '-')),
            str(qty),
            f"{float(item.get('tool_unit', 0)):.2f} zł",
            f"{t_tool:.2f} zł",
            str(item.get('coat_name', '-')),
            f"{float(item.get('coat_unit', 0)):.2f} zł",
            f"{t_coat:.2f} zł",
            uslugi_str,
            f"{t_extra:.2f} zł",
            str(item.get('notes', ''))
        ]

        for i, val in enumerate(row_data):
            row_cells[i].text = val
            p = row_cells[i].paragraphs[0]
            for run in p.runs:
                run.font.size = Pt(7)
                run.font.name = 'Arial'
            
            if idx % 2 == 0:
                set_cell_background(row_cells[i], "E2E8F0")

    doc.add_paragraph()

    grand_total = total_tool_val + total_coat_val + total_extra_val

    sum_table = doc.add_table(rows=1, cols=2)
    sum_table.autofit = False
    sum_table.columns[0].width = Inches(7.0)
    sum_table.columns[1].width = Inches(3.5)

    sum_cell = sum_table.cell(0, 1)
    set_cell_background(sum_cell, "F1F5F9")
    
    p_sum = sum_cell.paragraphs[0]
    p_sum.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    sum_text = (
        f"Suma sztuk: {total_qty} szt.\n"
        f"Wartość regeneracji: {total_tool_val:.2f} zł\n"
        f"Wartość powlekania: {total_coat_val:.2f} zł\n"
        f"Wartość usług dod.: {total_extra_val:.2f} zł\n"
    )
    r_sum = p_sum.add_run(sum_text)
    r_sum.font.size = Pt(8)
    r_sum.font.name = 'Arial'

    r_total = p_sum.add_run(f"RAZEM NETTO: {grand_total:.2f} zł")
    r_total.bold = True
    r_total.font.size = Pt(10)
    r_total.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    r_total.font.name = 'Arial'

    doc.save(output_docx_path)